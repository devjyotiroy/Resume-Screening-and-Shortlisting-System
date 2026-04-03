import pdfplumber
import docx
import re

def extract_text_from_pdf(path):
    text = ''
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # Extract regular text
            t = page.extract_text()
            if t:
                text += t + '\n'
            # Also extract tables (catches skill tables in resumes)
            for table in page.extract_tables():
                for row in table:
                    row_text = ' '.join(cell for cell in row if cell)
                    text += row_text + '\n'
    return text

def extract_text_from_docx(path):
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Extract table text too
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)

def clean_text(text):
    """Remove excessive whitespace and non-printable chars."""
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()

def extract_text(path):
    raw = ''
    if path.lower().endswith('.pdf'):
        raw = extract_text_from_pdf(path)
    elif path.lower().endswith('.docx'):
        raw = extract_text_from_docx(path)
    return clean_text(raw)
